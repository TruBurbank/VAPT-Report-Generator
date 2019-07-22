import os
from docx import Document
from docx.shared import Inches
from docx import section
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
import docx

class Print_document():

    def start_doc(self):
        self.document = Document()

    def reinitialize_doc(self):
        self.document = Document('Temp.docx')

    def initialize_doc(self):
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        style = self.document.styles['Heading 2']
        font1 = style.font
        font1.name = 'Times New Roman'
        font1.size = Pt(16)

        header = self.document.sections[0].header
        ht0=header.add_paragraph()
        kh=ht0.add_run()
        kh.add_picture('Pristine.png', width=Inches(2))
        kh.alignment = WD_ALIGN_PARAGRAPH.LEFT

        footer = self.document.sections[0].footer
        f = footer.add_paragraph('All Rights Reserved by Pristine InfoSolutions Pvt. Ltd.')
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f.style = self.document.styles['Normal']
        f.bold = True
        f.size = Pt(16)

    def setVname(self,Vname):
        
        self.document.add_heading('Vulnerability Name:', 2) 
        p = self.document.add_paragraph(Vname)
        p.style = self.document.styles['Normal']

    def setTitle(self):
        '''style1 = self.document.styles['Heading 3']
        font2 = style1.font
        font2.name = 'TimesNewRoman'
        font2.size = Pt(48)'''
        p = self.document.add_heading("NETWORK", 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #p.style1 = self.document.styles['Heading 3']
        '''obj_styles = self.document.styles
        obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(18)
        obj_font.name = 'Cambria (Headings)'''
        p1 = self.document.add_paragraph('VULNERABILITY ASSESSMENT & PENETRATION TESTING')
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        h1 = self.document.add_heading("REPORT",0 )
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h2 = self.document.add_heading("(INTERIM)",0)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER




    def setCName(self,cname):
        p = self.document.add_heading(cname, 2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1 = self.document.add_paragraph('The information contained within this report is considered proprietary and confidential to the '+ cname +' Inappropriate and unauthorized disclosure of this report or portions of it could result in significant damage or loss to the '+cname + ' This report should be distributed to individuals on a Need-to-Know basis only. Paper copies should be locked up when not in use. Electronic copies should be stored offline and protected.')
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def AuthorTable(self,author,classification,approach,manager,title,version):
        p = self.document.add_heading('REPORT DETAILS', 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        records = (
            (title, version, author,manager,classification,approach ),
        )

        table = self.document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Title'
        hdr_cells[1].text = 'Version'
        hdr_cells[2].text = 'Author'
        hdr_cells[3].text = 'Approved By'
        hdr_cells[4].text = 'Classification'
        hdr_cells[5].text = 'Approach'
        for qty, id, desc, app, clfc, apr in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc
            row_cells[3].text = app
            row_cells[4].text = clfc
            row_cells[5].text = apr

    def RecipientTable(self,name,title,company):
        p = self.document.add_heading('RECIPIENT', 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        records = (
            (name, title, company),
        )

        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Company'
        for qty, id, desc, in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc

    def VersionTable(self,version,date,author):
        p = self.document.add_heading('Version Control', 0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        records = (
            (version, date, author),
        )

        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Version'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Author'
        for qty, id, desc, in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc            

    def setSummary(self):
        h = self.document.add_heading('Executive Summary', 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1 = self.document.add_paragraph('The security assessments performed, follow a standard assessment methodology beginning with reconnaissance, vulnerability enumeration and penetration testing for validation. We performed these assessments with the least possible impact to the organization. This means our assessment tools have been throttled back as to not consume customer bandwidth. Our assessments are also done at a mutually agreeable time which is determined to be least impacting to the organization.')
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.document.add_page_break()
        h1 = self.document.add_heading('Summary',0)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = self.document.add_paragraph('Penetration testing is the technique for exploiting the vulnerabilities present on the target, which has been earlier disclosed while performing vulnerability assessment. It’s a way to check if the Applications and Infrastructure is secure from accessing or compromising servers by an external users or guest users. Systems can be vulnerable which can be misused by the attackers once known, so to avoid such scenarios a Network VAPT should be done on the Network for a safe hack proof environment. In this engagement, testers will be trying to perform the tests by Black Box testing.  ')
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.document.add_page_break()
        h2 = self.document.add_heading('Scope of work',0)
        p3 = self.document.add_paragraph('Network VAPT, or penetration testing, is an important task to be carried out by IT administrators. This is because of the rise in hacking attempts irrespective of the industry type. Attacks can happen from internally or externally with no or little knowledge ')
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        l1 = self.document.add_paragraph(style = 'List Bullet')
        l1r = l1.add_run('Data collection').bold = True
        style = self.document.styles['List Bullet']
        font1 = style.font
        font1.name = 'Times New Roman'
        font1.color.rgb = RGBColor(23, 54, 93)
        font1.size = Pt(18)
        l1.style = self.document.styles['List Bullet']
        p5 = self.document.add_paragraph('Various methods including Google search are used to get target system data. One can also use web page source code analysis technique to get more info about the system, software and plug-in versions. There are many free tools and services available in the market which can give information like database or table names, DB versions, software versions, hardware used and various third-party plugins used in the target system. ')
        
        l2 = self.document.add_paragraph(style = 'List Bullet')
        l2r = l2.add_run('Vulnerability Assessment ').bold = True
        l2.style = self.document.styles['List Bullet']
        p6 = self.document.add_paragraph('Based on the data collected in first step one can find the security weakness in the target system. This helps penetration testers to launch attacks using identified entry points in the system. ')
        
        l3 = self.document.add_paragraph(style = 'List Bullet')
        l3r = l3.add_run('Vulnerability Exploitation ').bold = True
        l3.style = self.document.styles['List Bullet']
        p7 = self.document.add_paragraph('This step requires special skills and techniques to launch attack on target system. Experienced penetration testers can use their skills to launch attack on the system. ')        
        
        l4 = self.document.add_paragraph(style = 'List Bullet')
        l4r = l4.add_run('Result analysis and report preparation  ').bold = True
        l4.style = self.document.styles['List Bullet']
        p8 = self.document.add_paragraph('After completion of penetration tests detailed reports are prepared for taking corrective actions. All identified vulnerabilities and recommended corrective methods are listed in these reports. You can customize vulnerability report format (MS Word or PDF) as per your organization needs. ')

        p9 = self.document.add_paragraph('In this Penetration Testing, the testers will be trying to perform the above mentioned tests by Black Box testing. Black Box testing is when the testers are not aware of the infrastructure, design and implementation of the product being tested.')
        self.document.add_page_break()

    def setVSeverity(self,severity):
        p = self.document.add_heading('Severity', 2)
        p.style = self.document.styles['Heading 2']
        p.bold = True
        p.size = Pt(16)
        p.name = 'TimesNewRoman'
        p = self.document.add_paragraph(severity)
        p.style = self.document.styles['Normal']

    def SetVdesc(self,VDesc):
        vuldesh = self.document.add_heading('Vulnerability Description:', 2)
        p = self.document.add_paragraph(VDesc)

    def setVurl(self,Vurl):
        self.document.add_heading('Vulnerable URL: ', 2)
        p = self.document.add_paragraph(Vurl)
        p.style = self.document.styles['Normal']

    def setVport(self,Vport):
        self.document.add_heading('Vulnerable Port: ', 2)
        p = self.document.add_paragraph(Vport)
        p.style = self.document.styles['Normal']

    def setImg(self,Img):
        self.document.add_heading('Proof of Concept: ',2)
        if (Img):
            lengthImg = len(Img[0]) 
            for i in range (0,lengthImg):
                self.document.add_picture(Img[0][i], width=Cm(15.95))

    def setImpact(self,VImpact):
        self.document.add_heading('Impact: ',2)
        p = self.document.add_paragraph(VImpact)
        p.style = self.document.styles['Normal']

    def setVremed(self,Vrem):
        self.document.add_heading('Remediation:', 2 )
        p = self.document.add_paragraph(Vrem)
        p.style = self.document.styles['Normal']

    def setConclusion(self,Conclusion):
        self.document.add_heading('Conclusion:', 2 )
        p = self.document.add_paragraph(Conclusion)
        p.style = self.document.styles['Normal']

    def pageBreak(self):
        self.document.add_page_break()
    
    def Savedoc(self,name):
        self.document.save(name[0] + '.docx')

    def Savereport(self):
        self.document.save('Temp.docx')
